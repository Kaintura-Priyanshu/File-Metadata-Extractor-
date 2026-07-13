from pathlib import Path
from typing import Tuple, Dict, Any, Optional

from src.parsers.base_parser import BaseParser
from src.core.metadata_types import DocumentMetadata


class DocumentParser(BaseParser):
    """Extract metadata from PDF, Word, Excel, PowerPoint, and plain-text files."""

    def parse(self, file_path: Path) -> Tuple[Optional[DocumentMetadata], Dict[str, Any]]:
        metadata = DocumentMetadata()
        raw_metadata: Dict[str, Any] = {}

        file_extension = file_path.suffix.lower()

        try:
            if file_extension == '.pdf':
                self._parse_pdf(file_path, metadata, raw_metadata)
            elif file_extension in ('.docx', '.doc'):
                self._parse_word(file_path, metadata, raw_metadata)
            elif file_extension in ('.xlsx', '.xls'):
                self._parse_excel(file_path, metadata, raw_metadata)
            elif file_extension in ('.pptx', '.ppt'):
                self._parse_powerpoint(file_path, metadata, raw_metadata)
            elif file_extension == '.txt':
                self._parse_text(file_path, metadata, raw_metadata)
        except Exception as e:
            print(f"Error parsing document metadata: {e}")

        return metadata, raw_metadata

    def _parse_pdf(self, file_path: Path, metadata: DocumentMetadata, raw_metadata: Dict) -> None:
        try:
            from pypdf import PdfReader

            reader = PdfReader(file_path)

            if reader.is_encrypted:
                try:
                    reader.decrypt('')
                except Exception:
                    print("PDF is encrypted; metadata may be incomplete")

            info = reader.metadata

            if info:
                # Mutate the caller's dict in place (do NOT reassign
                # raw_metadata here — that would just rebind the local
                # name and silently discard the data on return).
                raw_metadata.update({key: str(value) for key, value in info.items()})

                if info.author:
                    metadata.author = str(info.author)
                if info.title:
                    metadata.title = str(info.title)
                if info.subject:
                    metadata.subject = str(info.subject)
                if info.get('/Keywords'):
                    keywords = str(info['/Keywords']).split(',')
                    metadata.keywords = [k.strip() for k in keywords if k.strip()]
                if info.creator:
                    metadata.application_name = str(info.creator)
                if info.producer:
                    metadata.application_version = str(info.producer)
                if info.creation_date:
                    metadata.creation_date = info.creation_date
                if info.modification_date:
                    metadata.last_modified_date = info.modification_date

            metadata.page_count = len(reader.pages)

            text_chunks = []
            for page in reader.pages:
                try:
                    text_chunks.append(page.extract_text() or "")
                except Exception:
                    continue

            text_content = "".join(text_chunks)
            if text_content:
                metadata.word_count = len(text_content.split())
                metadata.character_count = len(text_content.replace(" ", ""))

        except ImportError:
            print("pypdf library not available for PDF parsing")
        except Exception as e:
            print(f"Error parsing PDF: {e}")

    def _parse_word(self, file_path: Path, metadata: DocumentMetadata, raw_metadata: Dict) -> None:
        try:
            from docx import Document

            doc = Document(file_path)
            core_props = doc.core_properties

            if core_props.author:
                metadata.author = core_props.author
            if core_props.title:
                metadata.title = core_props.title
            if core_props.subject:
                metadata.subject = core_props.subject
            if core_props.keywords:
                metadata.keywords = [k.strip() for k in core_props.keywords.split(',') if k.strip()]
            if core_props.created:
                metadata.creation_date = core_props.created
            if core_props.modified:
                metadata.last_modified_date = core_props.modified

            metadata.application_name = "Microsoft Word"

            text_chunks = [paragraph.text for paragraph in doc.paragraphs]
            text_content = "\n".join(text_chunks)

            if text_content:
                metadata.word_count = len(text_content.split())
                metadata.character_count = len(text_content.replace(" ", ""))

            metadata.page_count = len(doc.sections)

            raw_metadata.update({
                'paragraph_count': len(doc.paragraphs),
                'table_count': len(doc.tables),
            })

        except ImportError:
            print("python-docx library not available for Word parsing")
        except Exception as e:
            print(f"Error parsing Word document: {e}")

    def _parse_excel(self, file_path: Path, metadata: DocumentMetadata, raw_metadata: Dict) -> None:
        try:
            import openpyxl

            workbook = openpyxl.load_workbook(file_path, data_only=True, read_only=True)

            props = workbook.properties
            if props:
                if props.creator:
                    metadata.author = props.creator
                if props.title:
                    metadata.title = props.title
                if props.subject:
                    metadata.subject = props.subject
                if props.keywords:
                    keywords = props.keywords.split(',')
                    metadata.keywords = [k.strip() for k in keywords if k.strip()]
                if props.created:
                    metadata.creation_date = props.created
                if props.modified:
                    metadata.last_modified_date = props.modified

            metadata.application_name = "Microsoft Excel"

            total_cells = 0
            for sheet in workbook.worksheets:
                for row in sheet.iter_rows():
                    total_cells += len(row)

            metadata.word_count = total_cells
            raw_metadata['sheet_names'] = workbook.sheetnames
            raw_metadata['sheet_count'] = len(workbook.sheetnames)

            workbook.close()

        except ImportError:
            print("openpyxl library not available for Excel parsing")
        except Exception as e:
            print(f"Error parsing Excel document: {e}")

    def _parse_powerpoint(self, file_path: Path, metadata: DocumentMetadata, raw_metadata: Dict) -> None:
        try:
            from pptx import Presentation

            prs = Presentation(file_path)
            core_props = prs.core_properties

            if core_props.author:
                metadata.author = core_props.author
            if core_props.title:
                metadata.title = core_props.title
            if core_props.subject:
                metadata.subject = core_props.subject
            if core_props.keywords:
                keywords = core_props.keywords.split(',')
                metadata.keywords = [k.strip() for k in keywords if k.strip()]
            if core_props.created:
                metadata.creation_date = core_props.created
            if core_props.modified:
                metadata.last_modified_date = core_props.modified

            metadata.application_name = "Microsoft PowerPoint"
            metadata.page_count = len(prs.slides)

        except ImportError:
            print("python-pptx library not available for PowerPoint parsing")
        except Exception as e:
            print(f"Error parsing PowerPoint document: {e}")

    def _parse_text(self, file_path: Path, metadata: DocumentMetadata, raw_metadata: Dict) -> None:
        content = None

        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
        except UnicodeDecodeError:
            try:
                with open(file_path, 'r', encoding='latin-1') as f:
                    content = f.read()
                raw_metadata['encoding_fallback'] = 'latin-1'
            except Exception as e:
                print(f"Error reading text file: {e}")
                return
        except Exception as e:
            print(f"Error reading text file: {e}")
            return

        if content is None:
            return

        metadata.word_count = len(content.split())
        metadata.character_count = len(content.replace(" ", ""))
        metadata.page_count = max(1, (len(content) + 2000) // 2000)

        raw_metadata['content_length'] = len(content)
        raw_metadata['line_count'] = content.count('\n') + 1
